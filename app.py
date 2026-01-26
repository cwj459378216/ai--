import asyncio
import os
from pathlib import Path
from typing import Any, Optional

import pandas as pd
import pdfplumber
import requests
import streamlit as st
from docx import Document as DocxDocument
from pymilvus import MilvusClient

try:
    from dotenv import load_dotenv

    load_dotenv()
except Exception:
    pass

from llama_index.core import StorageContext, VectorStoreIndex
from llama_index.core.base.embeddings.base import BaseEmbedding
from llama_index.core.schema import Document
from llama_index.vector_stores.milvus import MilvusVectorStore


def run_coro(coro):
    """在 Streamlit 这种同步环境里安全跑 async。"""

    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


class BaiLianEmbedding(BaseEmbedding):
    model_name: str = "bailian"
    api_key: str
    endpoint: str = (
        "https://dashscope.aliyuncs.com/api/v1/services/embeddings/text-embedding/text-embedding"
    )
    model: str = "text-embedding-v2"

    def _get_embedding_with_type(self, text: str, text_type: str):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        data = {
            "model": self.model,
            "input": {"texts": [text], "text_type": text_type},
        }
        resp = requests.post(self.endpoint, json=data, headers=headers, timeout=60)
        resp.raise_for_status()

        try:
            payload = resp.json()
        except ValueError as exc:
            body_preview = (resp.text or "").strip()[:800]
            raise RuntimeError(
                f"Embedding API returned non-JSON response (status={resp.status_code}): {body_preview}"
            ) from exc

        try:
            return payload["output"]["embeddings"][0]["embedding"]
        except Exception as exc:
            raise RuntimeError(
                f"Unexpected embedding response shape: keys={list(payload.keys())} payload_preview={str(payload)[:800]}"
            ) from exc

    def _get_text_embedding(self, text: str):
        return self._get_embedding_with_type(text, "document")

    def _get_query_embedding(self, query: str):
        return self._get_embedding_with_type(query, "query")

    async def _aget_text_embedding(self, text: str):
        return await asyncio.to_thread(self._get_text_embedding, text)

    async def _aget_query_embedding(self, query: str):
        return await self._aget_text_embedding(query)


def dashscope_chat(*, api_key: str, model: str, system: str, user: str) -> str:
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    data: dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
    }

    resp = requests.post(url, json=data, headers=headers, timeout=120)
    if resp.status_code >= 400:
        body_preview = (resp.text or "").strip()[:1500]
        raise RuntimeError(f"LLM API error status={resp.status_code}: {body_preview}")

    payload = resp.json()
    return payload["choices"][0]["message"]["content"].strip()


def load_documents_from_dir(data_dir: str) -> list[Document]:
    data_dir_path = Path(data_dir)
    documents: list[Document] = []

    if not data_dir_path.exists():
        return documents

    for file in data_dir_path.iterdir():
        if file.is_dir():
            continue
        suffix = file.suffix.lower()

        if suffix == ".txt":
            text = file.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif suffix == ".docx":
            doc = DocxDocument(file)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif suffix == ".csv":
            df = pd.read_csv(file)
            text = "\n".join(df.astype(str).agg(" ".join, axis=1))
        else:
            continue

        documents.append(Document(text=text, doc_id=file.name))

    return documents


async def build_index(*, api_key: str, docs_dir: str, db_path: str, collection: str, dim: int) -> None:
    documents = load_documents_from_dir(docs_dir)
    if not documents:
        raise RuntimeError(f"docs 目录为空或不存在：{docs_dir}")

    milvus_path = Path(db_path)
    milvus_path.parent.mkdir(parents=True, exist_ok=True)

    milvus_client = MilvusClient(uri=str(milvus_path.resolve()))
    vector_store = MilvusVectorStore(
        client=milvus_client,
        collection_name=collection,
        dim=dim,
        overwrite=False,
    )

    embed_model = BaiLianEmbedding(api_key=api_key)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        embed_model=embed_model,
    )


async def query_rag(
    *,
    api_key: str,
    question: str,
    topk: int,
    llm_model: str,
    db_path: str,
    collection: str,
    dim: int,
) -> tuple[list[dict[str, Any]], str]:
    milvus_path = Path(db_path)
    milvus_path.parent.mkdir(parents=True, exist_ok=True)

    milvus_client = MilvusClient(uri=str(milvus_path.resolve()))
    vector_store = MilvusVectorStore(
        client=milvus_client,
        collection_name=collection,
        dim=dim,
        overwrite=False,
    )

    embed_model = BaiLianEmbedding(api_key=api_key)

    index = VectorStoreIndex.from_vector_store(vector_store=vector_store, embed_model=embed_model)
    retriever = index.as_retriever(similarity_top_k=topk)
    nodes = retriever.retrieve(question)

    results: list[dict[str, Any]] = []
    context_blocks: list[str] = []

    for i, node in enumerate(nodes, start=1):
        text = (node.get_content() or "").strip()
        score = getattr(node, "score", None)
        results.append(
            {
                "rank": i,
                "score": score,
                "preview": text[:800],
                "text": text,
            }
        )
        context_blocks.append(f"[{i}] {text}")

    if not results:
        return [], "未召回到任何内容。请先入库或检查 collection/dim 是否匹配。"

    context = "\n\n".join(context_blocks)
    user_prompt = (
        "请只基于【上下文】回答【问题】。如果资料不足，请明确说“不确定/资料不足”。\n\n"
        f"【上下文】\n{context}\n\n"
        f"【问题】\n{question}\n\n"
        "【回答】\n"
    )

    answer = dashscope_chat(
        api_key=api_key,
        model=llm_model,
        system="你是一个严谨的助手。",
        user=user_prompt,
    )

    return results, answer


st.set_page_config(page_title="本地 Milvus-lite RAG", layout="wide")
st.title("本地 Milvus-lite RAG 操作页")

with st.sidebar:
    st.header("配置")

    api_key_default = os.getenv("DASHSCOPE_API_KEY") or os.getenv("BAILIAN_API_KEY") or ""
    api_key = st.text_input("DASHSCOPE_API_KEY", value=api_key_default, type="password")

    llm_model = st.selectbox("生成模型", ["qwen-plus", "qwen-turbo", "qwen-max"], index=0)
    topk = st.slider("TopK", min_value=1, max_value=10, value=3)

    docs_dir = st.text_input("docs 目录", value="./docs")
    db_path = st.text_input("milvus.db 路径", value="./milvus_data/milvus.db")
    collection = st.text_input("collection", value="rag_demo")
    dim = st.number_input("向量维度（dim）", min_value=1, value=1536, step=1)

    st.caption("提示：建议把 key 写进 .env（不要提交到 git）。")

col1, col2 = st.columns([2, 3])

with col1:
    st.subheader("入库（可选）")
    st.write("如果你已经跑过 2_vector_to_milvus.py 并写入 Milvus，可跳过。")

    if st.button("一键入库/更新", type="primary"):
        if not api_key:
            st.error("请先在左侧填写 DASHSCOPE_API_KEY")
        else:
            with st.spinner("正在向量化并写入 Milvus-lite..."):
                try:
                    run_coro(
                        build_index(
                            api_key=api_key,
                            docs_dir=docs_dir,
                            db_path=db_path,
                            collection=collection,
                            dim=int(dim),
                        )
                    )
                    st.success("入库完成")
                except Exception as exc:
                    st.error(str(exc))

with col2:
    st.subheader("查询 + RAG")
    question = st.text_input("问题", value="买入日期是2026-01-20的股票")

    if st.button("查询", type="primary"):
        if not api_key:
            st.error("请先在左侧填写 DASHSCOPE_API_KEY")
        else:
            with st.spinner("检索中..."):
                try:
                    results, answer = run_coro(
                        query_rag(
                            api_key=api_key,
                            question=question,
                            topk=int(topk),
                            llm_model=llm_model,
                            db_path=db_path,
                            collection=collection,
                            dim=int(dim),
                        )
                    )

                    st.markdown("### RAG 回答")
                    st.write(answer)

                    st.markdown("### TopK 检索结果")
                    if results:
                        for item in results:
                            with st.expander(f"#{item['rank']} score={item['score']}", expanded=item["rank"] == 1):
                                st.text(item["preview"])
                    else:
                        st.info("未召回到内容")

                except Exception as exc:
                    st.error(str(exc))
