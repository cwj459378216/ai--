from __future__ import annotations

from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from llama_index.core.schema import Document

from .chunking import chunk_text_by_chars


SUPPORTED_SUFFIXES = {".txt", ".pdf", ".docx", ".csv"}


def load_documents_from_dir(
    data_dir: Path,
    *,
    chunk_size: int = 500,
    overlap: int = 50,
) -> list[Document]:
    data_dir = Path(data_dir)
    documents: list[Document] = []

    for file in data_dir.iterdir():
        if not file.is_file():
            continue
        if file.name.startswith("."):
            continue

        suffix = file.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            print(f"⚠️ 未支持的文件类型: {file.name}")
            continue

        if suffix == ".txt":
            text = file.read_text(encoding="utf-8")
            documents.extend(_chunk_text(file, text, chunk_size, overlap))
        elif suffix == ".pdf":
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            documents.extend(_chunk_text(file, text, chunk_size, overlap))
        elif suffix == ".docx":
            doc = DocxDocument(file)
            text = "\n".join(p.text for p in doc.paragraphs)
            documents.extend(_chunk_text(file, text, chunk_size, overlap))
        elif suffix == ".csv":
            df = pd.read_csv(file)
            text = "\n".join(df.astype(str).agg(" ".join, axis=1))
            documents.extend(_chunk_text(file, text, chunk_size, overlap))

    return documents


def list_files(data_dir: Path) -> list[Path]:
    data_dir = Path(data_dir)
    return [p for p in data_dir.iterdir() if p.is_file()]


def _chunk_text(
    file: Path,
    text: str,
    chunk_size: int,
    overlap: int,
) -> list[Document]:
    chunks = chunk_text_by_chars(text, chunk_size=chunk_size, overlap=overlap)
    docs: list[Document] = []
    for i, c in enumerate(chunks, start=1):
        docs.append(
            Document(
                text=c.text,
                doc_id=f"{file.name}#chunk{i}",
                metadata={
                    "source": file.name,
                    "chunk": i,
                    "start_char": c.start,
                    "end_char": c.end,
                },
            )
        )
    return docs
