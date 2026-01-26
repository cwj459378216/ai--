from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
from llama_index.core.schema import Document

def load_documents_from_dir(data_dir: str):
    data_dir = Path(data_dir)
    documents = []

    for file in data_dir.iterdir():
        if file.suffix.lower() == ".txt":
            text = file.read_text(encoding="utf-8")
            documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".pdf":
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".docx":
            doc = DocxDocument(file)
            text = "\n".join(p.text for p in doc.paragraphs)
            documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".csv":
            df = pd.read_csv(file)
            text = "\n".join(df.astype(str).agg(" ".join, axis=1))
            documents.append(Document(text=text, doc_id=file.name))
        else:
            print(f"⚠️ 未支持的文件类型: {file.name}")
    
    return documents

if __name__ == "__main__":
    docs = load_documents_from_dir("./docs")
    print(f"✅ 共加载 {len(docs)} 个文档")