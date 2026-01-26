import asyncio
import os
from pathlib import Path
import requests

from llama_index.core.schema import Document
from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.vector_stores.milvus import MilvusVectorStore
from llama_index.core.base.embeddings.base import BaseEmbedding
from pymilvus import MilvusClient

try:
    from dotenv import load_dotenv

    load_dotenv()
except Exception:
    pass

# -------------------------------
# 自定义 BaiLianEmbedding
# -------------------------------
class BaiLianEmbedding(BaseEmbedding):
    model_name: str = "bailian"
    api_key: str
    endpoint: str = "https://dashscope.aliyuncs.com/api/v1/services/embeddings/text-embedding/text-embedding"
    model: str = "text-embedding-v2"

    def _get_embedding_with_type(self, text: str, text_type: str):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        data = {
            "model": self.model,
            "input": {
                "texts": [text],
                "text_type": text_type,
            },
        }
        resp = requests.post(self.endpoint, json=data, headers=headers, timeout=60)
        resp.raise_for_status()

        try:
            payload = resp.json()
        except ValueError as exc:
            body_preview = (resp.text or "").strip()[:500]
            raise RuntimeError(
                f"Embedding API returned non-JSON response (status={resp.status_code}): {body_preview}"
            ) from exc

        try:
            return payload["output"]["embeddings"][0]["embedding"]
        except Exception as exc:
            raise RuntimeError(
                f"Unexpected embedding response shape: keys={list(payload.keys())} payload_preview={str(payload)[:500]}"
            ) from exc

    def _get_text_embedding(self, text: str):
        return self._get_embedding_with_type(text, "document")

    def _get_query_embedding(self, query: str):
        return self._get_embedding_with_type(query, "query")

    async def _aget_text_embedding(self, text: str):
        return await asyncio.to_thread(self._get_text_embedding, text)

    async def _aget_query_embedding(self, query: str):
        return await self._aget_text_embedding(query)

# -------------------------------
# 文档加载函数
# -------------------------------
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd

def load_documents_from_dir(data_dir: str):
    data_dir = Path(data_dir)
    documents = []

    for file in data_dir.iterdir():
        if file.suffix.lower() == ".txt":
            text = file.read_text(encoding="utf-8")
            documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".pdf":
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".docx":
            doc = DocxDocument(file)
            text = "\n".join(p.text for p in doc.paragraphs)
            documents.append(Document(text=text, doc_id=file.name))
        elif file.suffix.lower() == ".csv":
            df = pd.read_csv(file)
            text = "\n".join(df.astype(str).agg(" ".join, axis=1))
            documents.append(Document(text=text, doc_id=file.name))
        else:
            print(f"⚠️ 未支持的文件类型: {file.name}")
    
    return documents

# -------------------------------
# 主流程
# -------------------------------
async def main():
    # 读取文档
    documents = load_documents_from_dir("./docs")

    # Milvus-lite 本地路径
    milvus_path = Path("./milvus_data/milvus.db")
    milvus_path.parent.mkdir(parents=True, exist_ok=True)

    # Milvus 客户端
    milvus_client = MilvusClient(uri=str(milvus_path.resolve()))

    # Milvus 向量存储
    vector_store = MilvusVectorStore(
        client=milvus_client,
        collection_name="rag_demo",
        dim=1536,  # 百炼 embedding 维度
        overwrite=False
    )

    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("BAILIAN_API_KEY")
    if not api_key:
        raise SystemExit("请先设置环境变量 DASHSCOPE_API_KEY（或 BAILIAN_API_KEY）")

    # 自定义 embedding
    embed_model = BaiLianEmbedding(api_key=api_key)

    # LlamaIndex 存储上下文
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # 构建向量索引
    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        embed_model=embed_model
    )

    print("✅ 文档向量化完成并写入 Milvus-lite")

if __name__ == "__main__":
    asyncio.run(main())